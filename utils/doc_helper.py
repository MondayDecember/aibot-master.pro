import asyncio
import io
import logging

from aiogram import Bot
from aiogram.types import Document

from utils.texts import t

logger = logging.getLogger(__name__)

# Telegram Bot API refuses to download files bigger than 20 MB
MAX_FILE_SIZE = 20 * 1024 * 1024

TEXT_EXTENSIONS = (
    ".txt", ".md", ".csv", ".tsv", ".log", ".json", ".xml", ".yaml", ".yml",
    ".html", ".htm", ".ini", ".cfg", ".toml", ".py", ".js", ".ts", ".sh",
    ".sql", ".java", ".c", ".cpp", ".h", ".go", ".rs", ".rb", ".php",
)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    """Text from a Word .docx: paragraphs plus table cells, in order."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _decode_text(data: bytes) -> str:
    """
    Decode a text file of unknown encoding. Plain .decode('utf-8') mangled
    Russian ANSI (Windows-1251) files into gibberish; here we detect the
    encoding (charset-normalizer) and fall back through the usual Russian
    suspects before giving up on a lossy replace.
    """
    # BOM-marked UTF variants decode cleanly and unambiguously first
    for bom, enc in ((b"\xef\xbb\xbf", "utf-8-sig"),
                     (b"\xff\xfe", "utf-16"), (b"\xfe\xff", "utf-16")):
        if data.startswith(bom):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                break
    try:
        from charset_normalizer import from_bytes
        best = from_bytes(data).best()
        if best is not None:
            return str(best)
    except Exception as e:
        logger.debug(f"charset detection failed, falling back: {e}")
    # Manual cascade: strict UTF-8, then Windows-1251 (Russian ANSI), then
    # a never-fails latin-1 so we always return *something*.
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


async def extract_document_text(bot: Bot, document: Document):
    """
    Download a telegram document and extract plain text from it.
    Returns (text, None) on success or (None, user_facing_error) on failure.
    """
    name = (document.file_name or "").lower()
    if document.file_size and document.file_size > MAX_FILE_SIZE:
        return None, t("doc_too_large")

    is_pdf = name.endswith(".pdf")
    is_docx = name.endswith(".docx")
    if not is_pdf and not is_docx and not name.endswith(TEXT_EXTENSIONS):
        return None, t("doc_unsupported")

    try:
        file_info = await bot.get_file(document.file_id)
        buf = io.BytesIO()
        await bot.download_file(file_info.file_path, buf)
        data = buf.getvalue()
    except Exception as e:
        logger.error(f"Failed to download document: {e}")
        return None, t("doc_download_failed")

    try:
        # PDF/DOCX parsing is CPU-bound - keep it off the event loop
        if is_pdf:
            text = await asyncio.to_thread(_extract_pdf, data)
        elif is_docx:
            text = await asyncio.to_thread(_extract_docx, data)
        else:
            text = await asyncio.to_thread(_decode_text, data)
    except Exception as e:
        logger.error(f"Failed to extract text from document '{name}': {e}")
        return None, t("doc_unreadable")

    text = text.strip()
    if not text:
        return None, t("doc_no_text")
    return text, None
