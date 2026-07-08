import io

from utils.doc_helper import _decode_text, _extract_docx


# --- encoding detection ---

def test_decodes_utf8():
    assert _decode_text("Привет, мир!".encode("utf-8")) == "Привет, мир!"


def test_decodes_windows1251_ansi():
    # Russian ANSI - the case that used to come out as gibberish
    raw = "Привет, как дела?".encode("cp1251")
    assert _decode_text(raw) == "Привет, как дела?"


def test_decodes_utf16_with_bom():
    raw = "Тест кодировки".encode("utf-16")  # includes a BOM
    assert _decode_text(raw) == "Тест кодировки"


def test_decodes_utf8_bom():
    raw = b"\xef\xbb\xbf" + "abc".encode("utf-8")
    assert _decode_text(raw) == "abc"


def test_never_raises_on_binary_junk():
    # must always return *something*, never blow up
    assert isinstance(_decode_text(bytes(range(256))), str)


# --- docx extraction ---

def test_extract_docx():
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Первый абзац")
    doc.add_paragraph("Второй абзац")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Ячейка A"
    table.rows[0].cells[1].text = "Ячейка B"
    buf = io.BytesIO()
    doc.save(buf)

    text = _extract_docx(buf.getvalue())
    assert "Первый абзац" in text
    assert "Второй абзац" in text
    assert "Ячейка A" in text and "Ячейка B" in text
